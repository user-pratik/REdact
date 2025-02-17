
from docx import Document
import os
import tempfile

from services.addons.direct import apply_direct_redaction  # ✅ Uses apply_direct_redaction
from services.img import process_image  # ✅ Uses process_image

def process_docx_file(docx_file, sensitivity_level, redact_ocr, redact_meta, redact_face, redact_license_plate, redact_signature, redact_nsfw, is_document):
    """
    Process a .docx file by redacting text and blurring images.
    """
    doc = Document(docx_file)
    upload_folder = 'static/uploads'

    ### 🔴 1. Fixing TEXT Redaction ###
    for para in doc.paragraphs:
        original_text = para.text.strip()
        if original_text:
            redacted_text = apply_direct_redaction(original_text, sensitivity_level)  # ✅ Using direct.py
            para.clear()
            para.add_run(redacted_text)

    ### 🔵 2. Handling IMAGE Redaction ###
    image_replacements = {}

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_filename = os.path.join(upload_folder, os.path.basename(rel.target_ref))

            # Save the original image temporarily
            with open(image_filename, "wb") as f:
                f.write(image_data)

            # Process the image using img.py
            processed_image = process_image(image_filename, redact_ocr, redact_meta, redact_face, 
                                            redact_license_plate, redact_signature, redact_nsfw, 
                                            is_document, sensitivity_level)

            # Store the processed image for later replacement
            with open(processed_image, "rb") as img_file:
                image_replacements[rel.target_ref] = img_file.read()

    # Replace original images with processed images
    for rel in doc.part.rels.values():
        if rel.target_ref in image_replacements:
            rel.target_part._blob = image_replacements[rel.target_ref]

    ### 🟢 3. Save the Redacted Document ###
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)

    return temp_file.name
